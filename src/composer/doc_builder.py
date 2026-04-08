import re
import io
from pathlib import Path
from PIL import Image as PILImage
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_docx(markdown_path, frames, output_path, manual_images=None):
    if manual_images is None:
        manual_images = {}
    all_images = {**frames, **manual_images}
    with open(markdown_path, "r", encoding="utf-8") as f:
        markdown = f.read()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    current_section_id = None
    section_index = 0
    for line in markdown.split("\n"):
        line = line.rstrip()
        if line.startswith("# ") and not line.startswith("## "):
            heading = doc.add_heading(line[2:].strip(), level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = heading.runs[0] if heading.runs else heading.add_run()
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)  # Azul corporativo oscuro
            heading.paragraph_format.space_after = Pt(18)
            current_section_id = f"sec_{section_index:02d}"
            section_index += 1
            continue
        if line.startswith("## "):
            section_title = line[3:].strip()
            doc.add_heading(section_title, level=1)
            current_section_id = f"sec_{section_index:02d}"
            section_index += 1
            if current_section_id in all_images:
                img_path = all_images[current_section_id]
                if Path(img_path).exists() and Path(img_path).stat().st_size > 0:
                    try:
                        # Normalizar imagen con Pillow para garantizar compatibilidad con python-docx
                        buf = io.BytesIO()
                        PILImage.open(img_path).convert("RGB").save(buf, format="JPEG")
                        buf.seek(0)
                        doc.add_picture(buf, width=Inches(5.5))
                        caption = doc.add_paragraph()
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = caption.add_run(f"Captura de video - {section_title}")
                        run.font.size = Pt(9)
                        run.font.italic = True
                    except Exception as e:
                        print(f"  Imagen omitida ({img_path}): {e}")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if re.match(r'^\s*[\*\-]\s+', line):
            text = re.sub(r'^\s*[\*\-]\s+', '', line).strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            doc.add_paragraph(text, style="List Bullet")
            continue
        if line.strip():
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', line.strip())
            doc.add_paragraph(text)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Documento DOCX generado: {output_path}")
    return output_path
